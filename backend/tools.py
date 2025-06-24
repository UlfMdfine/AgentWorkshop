import base64
import os
from pathlib import Path
from typing import List, Optional

import matplotlib
import pandas as pd
import numpy as np
from crsutil.plot.PlottingFrameworkPlotly import PlottingFrameworkPlotly
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv
from langchain.agents import tool
from langchain_core.messages import HumanMessage
from langchain_core.tools import StructuredTool
from langchain_openai import AzureChatOpenAI
from riskmodels.calibration.migration_matrix import (
    matrix_weighted_bandwidth,
    migration_matrix_diff_year,
)
from riskmodels.metrics.distribution_metrics import calc_psi
from riskmodels.metrics.hhi import calc_hhi
from riskmodels.plot.plot_utils import plotly_migration_matrix, plotly_psi
from riskmodels.plot.reviewofestimates.plot_roe_utils import (
    plot_discriminatory_power_over_time,
    plot_roe_weighted_distributions,
)
from riskmodels.representativeness import calc_frequency
from crsutil.plot.building_blocks.save_plot import save_plt_plot
from crsutil.plot.plotly_functions.bar_plots_plotly import create_plotly_bar_plot

from backend.helpers import (
    clean_and_prepare_data,
    create_sample_func,
    extract_title_from_summary,
    get_image_paths,
)

OUTPUT_DIR = Path("./outputs")
OUTPUT_DIR.mkdir(exist_ok=True)

matplotlib.use("Agg")  # Headless plot backend
load_dotenv()


# Initialize the LLM
llm = AzureChatOpenAI(
    azure_deployment="gpt-4.1",
    api_version="2024-12-01-preview",
    temperature=0,
    max_tokens=None,
    timeout=None,
    max_retries=2,
)


# ------------------- Tools -------------------
def plot_roe_distribution_tool_wrapper(
    csv_file_path: str,
    var_name: str,
    plot_group_by: str = "",
    weights: str = ""
) -> str:
    """
    Handles CSV data for ROE plotting by creating weighted distributions. The resulting file name will have '_distribution_weighted_by_<weights>' in it.

    Parameters
    ----------
    csv_file_path : str
        Path to the CSV file containing the data to be plotted.
        The CSV should have the expected columns required for analysis.

    var_name : str
        Name of the variable in the DataFrame for which the distribution plot is desired.
        This corresponds to a column name in the CSV file.

    output_path : str
        Directory path where the output plot files will be saved.
        Ensure valid write permissions exist for the specified path.

    pfw_obj : PlottingFramework
        An instance of the PlottingFramework class used to configure plot settings.
        This object manages aesthetics and saving configurations for the plots.

    plot_group_by : str, optional
        Variable name for grouping plots. Usually corresponds to a categorical column in the data.
        When specified, the plot will show distributions per group (default is empty, implying no grouping).

    show_plots_notebook : bool, optional
        Flag to indicate if plots should be rendered within Jupyter Notebook environments.
        Default is True to showcase plots interactively within supported interfaces.

    weights : str, optional
        Specifies the column name used to weight distribution plots.
        Useful for metrics like exposure or expected loss; applies weighted counts if given (default is empty).

    Returns
    -------
    str
        Success message indicating completion of the plot creation process or an error message detailing issues.

    """

    try:
        # Create pfw_obj manually
        pfw_obj=PlottingFrameworkPlotly()
        pfw_obj.load_corporate_design_d_fine()
        pfw_obj.plot_params_dict['yaxis']['automargin'] =True   

        # Verify CSV file path
        data_path = Path(csv_file_path)
        if not data_path.exists():
            return f"Error: CSV file '{csv_file_path}' not found."

        # Read CSV and create DataFrame
        df_in = clean_and_prepare_data(csv_file_path, var_name, weights, plot_group_by)
        
        # Call the plot function
        output_path = OUTPUT_DIR

        y_label = "Count"
        if weights:
            y_label_tmp = weights.split("_")
            y_label_tmp2 = [word[0].upper() + word[1:] for word in y_label_tmp]
            y_label = " ".join(y_label_tmp2)

        if plot_group_by != "":
            counts = []
            bin_edges = []
            hue_arr = []
            unique_groups = df_in[plot_group_by].unique()
            plot_title = ""
            for hue in unique_groups:
                if weights != "":
                    lcounts, lbin_edges = np.histogram(
                        df_in.loc[df_in[plot_group_by] == hue, var_name],
                        bins = [1,2,3,4,5,6,7],
                        weights=df_in.loc[df_in[plot_group_by] == hue, weights],
                    )
                    plot_title = f"{var_name}_distribution_weighted_by_{y_label}"
                else:
                    lcounts, lbin_edges = np.histogram(
                        df_in.loc[df_in[plot_group_by] == hue, var_name]
                    )
                    plot_title = f"{var_name}_distribution"

                lhue = [hue] * len(lcounts)
                for counter in range(0, len(lcounts)):
                    counts.append(lcounts[counter])
                    bin_edges.append(lbin_edges[counter])
                    hue_arr.append(lhue[counter])

            df = pd.DataFrame({"xvals": bin_edges, "yvals": counts, plot_group_by: hue_arr})

            fig = create_plotly_bar_plot(
                data=df,
                x="xvals",
                y="yvals",
                pfw_obj=pfw_obj,
                x_label=var_name,
                y_label=y_label,
                title=f"Distribution of {var_name}",
                hue=plot_group_by,
            )
        else:
            if weights != "":
                lcounts, lbin_edges = np.histogram(
                    df_in[var_name],
                    weights=df_in[weights],
                )
                plot_title = f"{var_name}_distribution_weighted_by_{y_label}"
            else:
                lcounts, lbin_edges = np.histogram(df_in[var_name])
                plot_title = f"{var_name}_distribution"
            df = pd.DataFrame({"xvals": lbin_edges, "yvals": lcounts})

            fig = create_plotly_bar_plot(
                data=df,
                x="xvals",
                y="yvals",
                pfw_obj=pfw_obj,
                x_label=var_name,
                y_label=y_label,
                title=f"Distribution of {var_name}",
            )

        save_plt_plot(
            fig=fig,
            plot_title=plot_title,
            output_dir_path=output_path,
    )


        
        return f"ROE distribution plot successfully created. Figure: {fig}"
    except Exception as e:
        return f"Error during ROE distribution plotting: {e}"
    
# Create the tool using StructuredTool from function
plot_roe_distribution_tool = StructuredTool.from_function(
    func=plot_roe_distribution_tool_wrapper,
    description="Generates a weighted ROE distribution plot based on a variable in a given CSV file.",
    parse_docstring=True,
)


def plot_roe_discriminatory_power_wrapper(
    csv_file_path: str,
    plot_by_time_var: str,
    var_list: List[str],
    target: str,
    method: str = "somers_d",
) -> str:
    """
    Calculates and plots the discriminatory power (AUROC, Somers' D or Gini coefficient) 
    for the CSV data in the ROE.

    Parameters
    ----------
    csv_file_path : str
        Path to the CSV file containing the data to be plotted.
        The CSV should have the expected columns required for analysis.
    plot_by_time_var: str
        (time) variable over which the discriminatory power is plottet. This is the reporting date.
    var_list : List[str]
        List of column names of the variables for which the discriminatory power metric
        should be calculated. This is the pd (probability of default) or the lgd (loss given default).
    target : str
        The name of the target variable the variables should be compared against, e.g. the default flag.
    method: str
        Method of discriminatory power analysis to be performed: 'somers_d' (also refered to as Somers' D, Somers D or Somer's D),
        'gini', or 'spearman'.

    Returns
    -------
    str
        Success message indicating completion of the plot creation process or an error message detailing issues.

    """

    try:
        # Create pfw_obj manually
        pfw_obj = PlottingFrameworkPlotly()
        pfw_obj.load_corporate_design_d_fine()
        pfw_obj.plot_params_dict['yaxis']['automargin'] =True   
        
        # Verify CSV file path
        data_path = Path(csv_file_path)
        if not data_path.exists():
            return f"Error: CSV file '{csv_file_path}' not found."

        # Read CSV and create DataFrame
        df_in = pd.read_csv(csv_file_path)
        # Ensure columns are numeric and handle NaNs
        for lvar in var_list:
            df_in.dropna(subset=[lvar], inplace=True)
        # drop NaNs
        df_in.dropna(subset=[target], inplace=True)
        df_in.dropna(subset=[plot_by_time_var], inplace=True)
        df_in["sample"] = "ALL"
        # Call the plot function
        output_path = OUTPUT_DIR

        discriminatory_power_over_timevar, fig = plot_discriminatory_power_over_time(
            df_in=df_in,
            plot_by_time_var=plot_by_time_var,
            var_list=var_list,
            target=target,
            pfw_obj=pfw_obj,
            method=method,
            group_by_list=["sample"],
            output_file_path=output_path,
            show_plots_notebook=False,
            show_time_dep_plots_notebook=False,
            create_roc_curve=True,
            boot_iter=10,
            confidence_level=0.05,
            min_num_observations_for_bootstrapping=5,
        )

        os.remove(Path(output_path) / Path(f'{method}_for_{var_list[0]}_by_{plot_by_time_var}_and_sample.html'))
        os.remove(Path(output_path) / Path(f'{method}_for_{var_list[0]}_by_{plot_by_time_var}_and_sample.png'))
        
        return f"Discriminatory power plot successfully created. Figure: {fig}"
    except Exception as e:
        return f"Error during discriminatory power plotting: {e}"


# Create the tool using StructuredTool from function
plot_roe_discriminatory_power_tool = StructuredTool.from_function(
    func=plot_roe_discriminatory_power_wrapper,
    description="Plots the discriminatory power based on a variable in a given CSV file.",
    parse_docstring=True,
)


def calc_migration_matrix_wrapper(
    csv_file_path: str,
    given_year: int,
    unique_id: str,
    rating_class_column: str,
) -> str:
    """
    Calculate the migration matrix of a given dataframe from given_year-1 to given_year.

    If a id/customer/facility doesn't have an entry at given_year-1 the transition is attributed to the rating class 0.

    Parameters
    ----------
    csv_file_path : str
        Path to the CSV file containing the data for which the migration matrix is calculated.
        The CSV should have the expected columns required for analysis.
    given_year: bool
        The year to which the migration is calculated.
    unique_id: bool
        The name of the existing column containing the unique identifier in the dataset.
    rating_class_column: bool
        The name of the existing column containing the rating class in the dataset.

    Returns
    -------
    str
        Success message indicating completion of the plot creation process or an error message detailing issues.

    """

    try:
        # Verify CSV file path
        data_path = Path(csv_file_path)
        if not data_path.exists():
            return f"Error: CSV file '{csv_file_path}' not found."

        # Read CSV and create DataFrame
        df_in = pd.read_csv(csv_file_path)
        df_in["reporting_date"] = pd.to_datetime(df_in["reporting_date"])
        df_in["year"] = df_in["reporting_date"].dt.year

        number_of_rating_grades = len(df_in["rating_class"].unique())
        # Call the plot function
        output_path = OUTPUT_DIR

        df_migration, df_migration_cleaned = migration_matrix_diff_year(
            df_in=df_in,
            given_year=given_year,
            unique_id=unique_id,
            reporting_date_year_var="year",
            rating_class_column=rating_class_column,
            number_of_rating_grades=number_of_rating_grades,
        )

        smpl_filename = f"migration_matrix.csv"
        output_path = OUTPUT_DIR / smpl_filename
        df_migration_cleaned.to_csv(output_path, index=False)

        return f"Migration matrix successfully created. Saved to: {output_path}"
    except Exception as e:
        return f"Error during calculation of migration matrix: {e}"


# Create the tool using StructuredTool from function
calc_migration_matrix_tool = StructuredTool.from_function(
    func=calc_migration_matrix_wrapper,
    description="Calculates a migration matrix based on a variable in a given CSV file.",
    parse_docstring=True,
)


def plot_migration_matrix_wrapper(
    csv_file_path: str,
    title: Optional[str],
    x_label: Optional[str],
    y_label: Optional[str],
) -> str:
    """
    Plot a migration matrix.

    Parameters
    ----------
    csv_file_path : str
        Path to the CSV file containing the data to be plotted.
        The CSV should have the data for the migration matrix that is plotted.
    title : Optional[str]
         The title of the plot.
    x_label : Optional[str]
        The x-label of the plot.
    y_label : Optional[str]
        The y-label of the plot.

    Returns
    -------
    str
        Success message indicating completion of the plot creation process or an error message detailing issues.

    """

    try:
        # Create pfw_obj manually
        pfw_obj = PlottingFrameworkPlotly()
        pfw_obj.load_corporate_design_d_fine()
        pfw_obj.plot_params_dict['yaxis']['automargin'] =True   

        # Verify CSV file path
        data_path = Path(csv_file_path)
        if not data_path.exists():
            return f"Error: CSV file '{csv_file_path}' not found."

        # Read CSV and create DataFrame
        df_in = pd.read_csv(csv_file_path)
        df_in.index = df_in.index + 1
        df_in.columns = pd.RangeIndex(start=1, stop=6, step=1)

        # Call the plot function
        output_path = OUTPUT_DIR
        plot_file_name = "migration_matrix_plot"

        fig = plotly_migration_matrix(
            df_in=df_in,
            title=title,
            plot_file_name=plot_file_name,
            pfw_obj=pfw_obj,
            x_label=x_label,
            y_label=y_label,
            output_dir_path=output_path,
            show_plots_notebook=False,
        )

        return f"Migration matrix plot successfully created. Figure: {fig}"
    except Exception as e:
        return f"Error during migration matrix plotting: {e}"


# Create the tool using StructuredTool from function
plot_migration_matrix_tool = StructuredTool.from_function(
    func=plot_migration_matrix_wrapper,
    description="Plots a migration matrix based on a CSV file.",
    parse_docstring=True,
)


def calc_matrix_weighted_bandwidth_wrapper(
    csv_file_path: str,
) -> str:
    """
    Calculate the upper and lower weighted bandwidth of the migration matrix.

    This implementation aligns with the Equations presented in "Instructions for reporting the validation results of
    internal models IRB Pillar I models for credit risk". Please note that we used the absolute migrations from the
    input matrix such that we did not need to multiply the equations with the numbers of customers in the original
    rating grades.

    Parameters
    ----------
    csv_file_path : str
        Path to the CSV file containing the matrix for which the bandwidth is calculated.

    Returns
    -------
    str
        Success message indicating completion of the plot creation process or an error message detailing issues.

    """
    try:
        # Verify CSV file path
        data_path = Path(csv_file_path)
        if not data_path.exists():
            return f"Error: CSV file '{csv_file_path}' not found."

        # Read CSV and create DataFrame
        df_in = pd.read_csv(csv_file_path)

        mwb_upper_normalized, mwb_lower_normalized = matrix_weighted_bandwidth(
            input_matrix=df_in
        )

        return f"calculated matrix bandwidth. Upper: {mwb_upper_normalized}; Lower: {mwb_lower_normalized}"
    except Exception as e:
        return f"Error during calculation of matrix bandwidth: {e}"


# Create the tool using StructuredTool from function
calc_matrix_weighted_bandwidth_tool = StructuredTool.from_function(
    func=calc_matrix_weighted_bandwidth_wrapper,
    description="Calculates a matrix bandwidth based on a matrix in a given CSV file.",
    parse_docstring=True,
)


def calc_psi_wrapper(
    csv_file_path: str,
    reporting_date_var: str,
    rating_class_column: str,
    date_to_split_samples: str,
) -> str:
    """
    Creates two samples based on the provided date and calculates the PSI of the rating classes of these samples.

    Parameters
    ----------
    csv_file_path : str
        Path to the CSV file containing the data for which the PSI is calculated.
        The CSV should have the expected columns required for analysis.
    reporting_date_var: bool
        The name of the existing column containing the reporting date.
    rating_class_column: bool
        The name of the existing column containing the rating class in df_in.
    date_to_split_samples: str
        The date on which the dataset is split into "old" and "new" subsets.

    Returns
    -------
    str
        Success message indicating completion of the plot creation process or an error message detailing issues.

    """

    try:
        # Verify CSV file path
        data_path = Path(csv_file_path)
        if not data_path.exists():
            return f"Error: CSV file '{csv_file_path}' not found."

        sample_dict = create_sample_func(
            file_path=csv_file_path,
            file_name=Path(csv_file_path).stem,
            date_column_name=reporting_date_var,
            date_for_split=date_to_split_samples,
        )

        if sample_dict["sample_new"][rating_class_column].shape[0] == 0:
            raise ValueError(
                f"The dataset does not contain any values after the given date used to split the dataset. "
                f"The date used to split was {date_to_split_samples}."
                # ToDo: I would remove the following addition for the last task.
                # f"Remark: You can use the tool to show the available values in the {reporting_date_var}."
            )
        elif sample_dict["sample_old"][rating_class_column].shape[0] == 0:
            raise ValueError(
                f"The dataset does not contain any values before the given date used to split the dataset. "
                f"The date used to split was {date_to_split_samples}. Please use another date."
                # f"Remark: You can use the tool to show the available values in the {reporting_date_var}."
            )

        freq_df = calc_frequency(
            distribution_1=sample_dict["sample_new"][rating_class_column],
            distribution_2=sample_dict["sample_old"][rating_class_column],
        )
        freq_df = pd.DataFrame(freq_df)
        freq_df.rename(
            columns={
                "extended_unique_counts_in_dis_1": "unique_counts_in_new_sample",
                "extended_unique_counts_in_dis_2": "unique_counts_in_old_sample",
            },
            inplace=True,
        )

        smpl_filename = f"rating_frequency_samples.csv"
        output_path = OUTPUT_DIR / smpl_filename
        freq_df.to_csv(output_path, index=False)

        dist1, dist2, psi_value = calc_psi(
            df_in=freq_df,
            c_abs_freq_1="unique_counts_in_new_sample",
            c_abs_freq_2="unique_counts_in_old_sample",
            risk_factor="Rating class",
        )

        return f"PSI successfully calculated as: {psi_value}. Frequency dataframes saved to: {output_path}"
    except Exception as e:
        return f"Error during calculation of PSI: {e}"


# Create the tool using StructuredTool from function
calc_psi_tool = StructuredTool.from_function(
    func=calc_psi_wrapper,
    description="Creates two samples based on the provided date and calculates the PSI of the rating classes of \n"
    "these samples for a given CSV file.",
    parse_docstring=True,
)


def plot_psi_wrapper(
    csv_file_path: str,
) -> str:
    """
    Plot the results of the psi calculation.

    Parameters
    ----------
    csv_file_path : str
        Path to the CSV file containing the data from the frequency analysis. This dataframe is build and calculated in
        calc_plotly_wrapper function.

    Returns
    -------
    str
        Success message indicating completion of the plot creation process or an error message detailing issues.

    """

    try:
        # Create pfw_obj manually
        pfw_obj = PlottingFrameworkPlotly()
        pfw_obj.load_corporate_design_d_fine()
        pfw_obj.plot_params_dict['yaxis']['automargin'] =True   

        # Verify CSV file path
        data_path = Path(csv_file_path)
        if not data_path.exists():
            return f"Error: CSV file '{csv_file_path}' not found."

        # Read CSV and create DataFrame
        df_freq = pd.read_csv(csv_file_path)

        # Call the plot function
        output_path = OUTPUT_DIR

        fig = plotly_psi(
            df_in=df_freq,
            c_class_interval="merged_unique",
            c_abs_freq_1="unique_counts_in_new_sample",
            c_abs_freq_2="unique_counts_in_old_sample",
            pfw_obj=pfw_obj,
            output_dir_path=str(output_path),
            show_plots_notebook=False,
        )

        return f"PSI plot successfully created. Figure: {fig}"
    except Exception as e:
        return f"Error during PSI plotting: {e}"


# Create the tool using StructuredTool from function
plot_psi_tool = StructuredTool.from_function(
    func=plot_psi_wrapper,
    description="Plots the results of the PSI analysis.",
    parse_docstring=True,
)


def calc_hhi_wrapper(
    csv_file_path: str,
    reporting_date_var: str,
    rating_class_column: str,
    date_to_split_samples: str,
) -> str:
    """
    Creates two samples based on the provided date and calculates the HHI's of the rating classes of these samples.

    Parameters
    ----------
    csv_file_path : str
        Path to the CSV file containing the data for which the PSI is calculated.
        The CSV should have the expected columns required for analysis.
    reporting_date_var: bool
        The name of the existing column containing the reporting date.
    rating_class_column: bool
        The name of the existing column containing the rating class in df_in.
    date_to_split_samples: str
        The date on which the dataset is split into "old" and "new" subsets.

    Returns
    -------
    str
        Success message indicating completion of the plot creation process or an error message detailing issues.

    """

    try:
        # Verify CSV file path
        data_path = Path(csv_file_path)
        if not data_path.exists():
            return f"Error: CSV file '{csv_file_path}' not found."

        sample_dict = create_sample_func(
            file_path=csv_file_path,
            file_name=Path(csv_file_path).stem,
            date_column_name=reporting_date_var,
            date_for_split=date_to_split_samples,
        )

        num_rating_classes = sample_dict["sample_new"][rating_class_column].nunique()
        rel_freq_dataset_new = sample_dict["sample_new"][
            rating_class_column
        ].value_counts(normalize=True)
        rel_freq_dataset_old = sample_dict["sample_old"][
            rating_class_column
        ].value_counts(normalize=True)

        if sample_dict["sample_new"][rating_class_column].shape[0] == 0:
            raise ValueError(
                f"The dataset does not contain any values after the given date used to split the dataset."
                f"The date used to split was {date_to_split_samples}."
                f"Remark: You can use the tool to show the available values in the {reporting_date_var}."
            )
        elif sample_dict["sample_old"][rating_class_column].shape[0] == 0:
            raise ValueError(
                f"The dataset does not contain any values before the given date used to split the dataset."
                f"The date used to split was {date_to_split_samples}. Please use another date."
                f"Remark: You can use the tool to show the available values in the {reporting_date_var}."
            )

        hhi_eu_new = calc_hhi(
            k_num_grade_non_defaulted=num_rating_classes,
            rel_freq_grade_start_obs=rel_freq_dataset_new,
        )
        hhi_eu_old = calc_hhi(
            k_num_grade_non_defaulted=num_rating_classes,
            rel_freq_grade_start_obs=rel_freq_dataset_old,
        )

        return (
            f"HHI's successfully calculated as:\n"
            f" - For sample 'old': {hhi_eu_old}\n"
            f" - For sample 'new': {hhi_eu_new}\n"
        )
    except Exception as e:
        return f"Error during calculation of HHI's: {e}"


# Create the tool using StructuredTool from function
calc_hhi_tool = StructuredTool.from_function(
    func=calc_hhi_wrapper,
    description="Creates two samples based on the provided date and calculates the HHI of the rating classes of \n"
    "these samples for a given CSV file.",
    parse_docstring=True,
)


def calc_unique_values_wrapper(
    csv_file_path: str,
    column_name: str,
) -> str:
    """
    Determines the unique values for a given column in the provided csv file.

    Parameters
    ----------
    csv_file_path : str
        Path to the CSV file containing the data.
        The CSV should have the expected columns required for analysis.
    column_name: str
        The name of the existing column.

    Returns
    -------
    str
        Success message indicating the unique values in the given column or an error message detailing issues.

    """

    try:
        # Verify CSV file path
        data_path = Path(csv_file_path)
        if not data_path.exists():
            return f"Error: CSV file '{csv_file_path}' not found."

        df_in = pd.read_csv(csv_file_path)
        unique_value_list = df_in[column_name].unique()

        return (
            f"Available values in colum {column_name} determined as:\n"
            f"{unique_value_list}"
        )
    except Exception as e:
        return f"Error during calculation of unique values: {e}"


# Create the tool using StructuredTool from function
calc_unique_values_tool = StructuredTool.from_function(
    func=calc_unique_values_wrapper,
    description="Determines the unique values of the given column in the dataset.",
    parse_docstring=True,
)


@tool
def summarize_chart(output_dir: Path) -> List[str]:
    """Summarize the content of all chart images in the specified directory."""
    summaries = []
    try:
        image_paths = get_image_paths(output_dir)
        if not image_paths:
            return ["Error: No images found in the specified directory."]

        for chart_image_path in image_paths:
            print(f"Summarizing chart: {chart_image_path}")
            if not os.path.exists(chart_image_path):
                summaries.append(f"Error: Chart image '{chart_image_path}' not found.")
                continue

            with open(chart_image_path, "rb") as f:
                image_base64 = base64.b64encode(f.read()).decode("utf-8")

            message = HumanMessage(
                content=[
                    {
                        "type": "text",
                        "text": f"Describe the content of the {chart_image_path} chart. "
                        "Focus on what stands out or is surprising. "
                        "Avoid obvious descriptions and instead infer meaningful patterns or anomalies. "
                        "Do not state that the summary is generated. "
                        "Provide a concise, insight-oriented summary suitable for including in a professional report."
                        "Use clear formatting (e.g., paragraphs) to make it easy to copy into a readable Word report.",
                    },
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{image_base64}"},
                    },
                ]
            )
            response = llm.invoke([message])
            summaries.append(response.content)

        return summaries

    except Exception as e:
        return [f"Error summarizing charts: {e}"]


@tool
def create_word_document(summaries: List[str], output_dir: Path) -> str:
    """Generate a Word document containing summaries and chart images using contextual matching."""
    try:
        doc = Document()
        doc.add_heading("AI-Generated Report", level=1)

        image_paths = get_image_paths(output_dir)
        print(f"Image paths: {image_paths}")  # Debugging

        for summary in summaries:
            title = extract_title_from_summary(summary)
            doc.add_heading(title, level=2)
            doc.add_paragraph(summary)
            doc.add_paragraph("")  # Add separation between sections

            # Extract keywords using the previously implemented function
            standard_keyword = title.replace(" Chart Summary", "")
            print(f"Standard keyword: {standard_keyword}")  # Debugging

            # Matching logic based on keyword presence in filenames
            matched_image_path = None
            for image_path in image_paths:
                # Allow for fuzzy matching based on presence of keywords in the filename
                if any(part.lower().replace(" ", "_") in image_path.stem.lower() for part in standard_keyword.split()):
                    matched_image_path = image_path
                    print(f"Matched image: {matched_image_path}")  # Debugging
                    break

            if matched_image_path:
                doc.add_heading(title.replace("Summary", "").strip(), level=2)
                with open(matched_image_path, "rb") as img_file:
                    doc.add_picture(img_file, width=Inches(5))
            
            # Explicit handling for special cases like bandwidth and HHI
            # Assume these are text summaries and ensure they're appended accordingly

        output_path = output_dir / "AI_Agent_Summary_Final.docx"
        doc.save(output_path)
        return f"Document saved as {output_path}."
    except Exception as e:
        return f"Error creating document: {e}"
